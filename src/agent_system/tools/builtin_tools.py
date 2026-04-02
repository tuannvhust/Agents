"""Built-in tools available to all agents out of the box.

Tools provided:
  - web_search       — DuckDuckGo web search (no API key required)
  - calculate        — Safe arithmetic / math expression evaluator
  - fetch_url        — HTTP GET a URL and return its text content
  - read_file        — Read a file from MinIO (or local path as fallback)
  - write_file       — Write text content to MinIO
  - create_word_file — Create a .docx file and save to MinIO
  - list_files       — List files in MinIO under an optional prefix
  - get_datetime     — Return the current UTC date and time
  - summarise_text   — Truncate long text to a readable excerpt
"""

from __future__ import annotations

import ast
import io
import logging
import math
import operator
import re
from datetime import datetime, timezone
from typing import Any

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from langchain_core.tools import tool

logger = logging.getLogger(__name__)


# ── Web Search ────────────────────────────────────────────────────────────────

@tool
def web_search(query: str, max_results: int = 5) -> str:
    """Search the web and return a list of results.

    Args:
        query: The search query string.
        max_results: Maximum number of results to return (default 5, max 10).

    Returns:
        Numbered list of search results with title, URL, and snippet.
    """
    try:
        from ddgs import DDGS

        max_results = min(max(1, max_results), 10)
        results = []
        with DDGS() as ddgs:
            # ddgs is the maintained successor of duckduckgo-search.
            for i, r in enumerate(ddgs.text(query, max_results=max_results), 1):
                title = r.get("title", "No title")
                href = r.get("href", "") or r.get("url", "")
                body = r.get("body", "") or r.get("snippet", "")
                results.append(f"{i}. **{title}**\n   URL: {href}\n   {body}")

        if not results:
            return "No results found."
        return "\n\n".join(results)

    except ImportError:
        return (
            "Web search dependency is not available. "
            "Install it with: pip install ddgs"
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("web_search failed: %s", exc)
        return f"Search error: {exc}"


# ── Calculator ────────────────────────────────────────────────────────────────

# Allowed AST node types for the safe evaluator
_SAFE_NODES = {
    ast.Expression, ast.BinOp, ast.UnaryOp, ast.Constant,
    ast.Add, ast.Sub, ast.Mult, ast.Div, ast.Mod, ast.Pow,
    ast.FloorDiv, ast.USub, ast.UAdd, ast.Call, ast.Name, ast.Load,
}

_SAFE_FUNCTIONS: dict[str, Any] = {
    "abs": abs, "round": round, "min": min, "max": max,
    "sqrt": math.sqrt, "log": math.log, "log2": math.log2,
    "log10": math.log10, "exp": math.exp, "ceil": math.ceil,
    "floor": math.floor, "sin": math.sin, "cos": math.cos,
    "tan": math.tan, "pi": math.pi, "e": math.e,
    "pow": math.pow, "factorial": math.factorial,
    "gcd": math.gcd, "hypot": math.hypot,
}


def _safe_eval(expression: str) -> float:
    """Evaluate a math expression using a whitelist-only AST evaluator."""
    tree = ast.parse(expression.strip(), mode="eval")

    for node in ast.walk(tree):
        if type(node) not in _SAFE_NODES:
            raise ValueError(f"Unsafe operation: {type(node).__name__}")
        if isinstance(node, ast.Name) and node.id not in _SAFE_FUNCTIONS:
            raise ValueError(f"Unknown name: {node.id!r}")
        if isinstance(node, ast.Call):
            if not isinstance(node.func, ast.Name):
                raise ValueError("Only simple function calls are allowed.")
            if node.func.id not in _SAFE_FUNCTIONS:
                raise ValueError(f"Unknown function: {node.func.id!r}")

    return eval(compile(tree, "<expr>", "eval"), {"__builtins__": {}}, _SAFE_FUNCTIONS)  # noqa: S307


@tool
def calculate(expression: str) -> str:
    """Evaluate a mathematical expression safely.

    Supports: +, -, *, /, //, %, ** and functions: abs, round, sqrt, log,
    log2, log10, exp, ceil, floor, sin, cos, tan, pow, factorial, gcd,
    hypot, min, max. Constants: pi, e.

    Args:
        expression: A math expression string, e.g. "sqrt(2) * pi" or "2**10".

    Returns:
        The numeric result as a string.
    """
    try:
        result = _safe_eval(expression)
        # Return int representation when the result is a whole number
        if isinstance(result, float) and result.is_integer():
            return str(int(result))
        return str(result)
    except ZeroDivisionError:
        return "Error: division by zero."
    except (ValueError, SyntaxError) as exc:
        return f"Error evaluating expression: {exc}"
    except Exception as exc:  # noqa: BLE001
        return f"Unexpected error: {exc}"


# ── HTTP Fetch ────────────────────────────────────────────────────────────────

@tool
def fetch_url(url: str, timeout: int = 15) -> str:
    """Fetch the text content of a URL via HTTP GET.

    Args:
        url: The full URL to fetch (must start with http:// or https://).
        timeout: Request timeout in seconds (default 15).

    Returns:
        The response body as plain text (first 8 000 characters).
    """
    if not url.startswith(("http://", "https://")):
        return "Error: URL must start with http:// or https://"
    try:
        headers = {"User-Agent": "AgentSystem/1.0 (research bot)"}
        with httpx.Client(timeout=timeout, follow_redirects=True) as client:
            resp = client.get(url, headers=headers)
            resp.raise_for_status()
            text = resp.text
            if len(text) > 8000:
                text = text[:8000] + f"\n\n[... truncated — {len(resp.text)} chars total]"
            return text
    except httpx.HTTPStatusError as exc:
        return f"HTTP error {exc.response.status_code}: {exc}"
    except httpx.RequestError as exc:
        return f"Request error: {exc}"


# ── File I/O (MinIO-backed) ───────────────────────────────────────────────────

def _get_minio():
    """Return a MinIOClient, or None if MinIO is not configured."""
    try:
        from agent_system.storage.minio_client import MinIOClient
        return MinIOClient()
    except Exception as exc:  # noqa: BLE001
        logger.warning("MinIO unavailable for file tool: %s", exc)
        return None


@tool
async def write_file(path: str, content: str) -> str:
    """Write text content to a file in MinIO object storage.

    Args:
        path: Object path / key inside the bucket, e.g. "reports/summary.txt".
        content: The text content to write.

    Returns:
        Confirmation message with the stored object path.
    """
    client = _get_minio()
    if client is None:
        return "Error: MinIO storage is not available."
    try:
        encoded = content.encode("utf-8")
        client.upload_bytes(path, encoded, content_type="text/plain")

        # Log to file_artifacts so every write is traceable by run_id / agent
        await _log_file_artifact(path, len(encoded), "text/plain")

        return f"File written successfully: {path}"
    except Exception as exc:  # noqa: BLE001
        return f"Error writing file: {exc}"


@tool
async def create_word_file(
    path: str,
    content: str,
    title: str = "",
    format: str = "plain",
    header_text: str = "",
    footer_text: str = "",
    include_toc: bool = False,
    image_sources: str = "",
) -> str:
    """Create a .docx file and upload it to MinIO.

    Args:
        path: MinIO object path ending with .docx, e.g. "reports/summary.docx".
        content: Main body content.
        title: Optional heading inserted at the top of the document.
        format: "plain" or "markdown".
        header_text: Optional header text for every page.
        footer_text: Optional footer text for every page.
        include_toc: Insert a table of contents field near the top.
        image_sources: Newline/comma separated list of image sources.
                       Supports MinIO object keys and http(s) URLs.

    Returns:
        Confirmation message with the stored object path.
    """
    if not path.lower().endswith(".docx"):
        return "Error: path must end with '.docx'."
    fmt = format.strip().lower()
    if fmt not in {"plain", "markdown"}:
        return "Error: format must be 'plain' or 'markdown'."

    client = _get_minio()
    if client is None:
        return "Error: MinIO storage is not available."

    try:
        doc = Document()
        _set_header_footer(doc, header_text=header_text, footer_text=footer_text)
        if title.strip():
            doc.add_heading(title.strip(), level=1)
        if include_toc:
            _add_table_of_contents(doc)
        if fmt == "markdown":
            await _add_markdown_to_doc(doc, content)
        else:
            _add_plain_text_to_doc(doc, content)
        await _append_images_to_doc(doc, image_sources=image_sources)

        buffer = io.BytesIO()
        doc.save(buffer)
        data = buffer.getvalue()

        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        client.upload_bytes(path, data, content_type=mime)
        await _log_file_artifact(file_path=path, file_size=len(data), content_type=mime)
        return f"Word file created successfully: {path}"
    except Exception as exc:  # noqa: BLE001
        logger.error("create_word_file failed: %s", exc)
        return f"Error creating Word file: {exc}"


def _add_plain_text_to_doc(doc: Document, content: str) -> None:
    chunks = [part.strip() for part in content.split("\n\n") if part.strip()]
    if not chunks:
        doc.add_paragraph("")
        return
    for paragraph_text in chunks:
        doc.add_paragraph(paragraph_text)


async def _add_markdown_to_doc(doc: Document, content: str) -> None:
    lines = content.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # Fenced code block: ``` ... ```
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                _add_code_block(doc, "\n".join(code_lines))
                in_code_block = False
                code_lines = []
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Headings: # ... to ###### ...
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            if 1 <= level <= 6 and stripped[level:].startswith(" "):
                doc.add_heading(stripped[level:].strip(), level=min(level, 4))
                i += 1
                continue

        # Tables: line with pipes followed by separator row
        if "|" in stripped and i + 1 < len(lines):
            sep = lines[i + 1].strip()
            if _is_markdown_table_separator(sep):
                headers = _split_md_row(stripped)
                aligns = _parse_md_alignments(sep)
                table_rows: list[list[str]] = []
                i += 2
                while i < len(lines):
                    row_line = lines[i].strip()
                    if not row_line or "|" not in row_line:
                        break
                    table_rows.append(_split_md_row(row_line))
                    i += 1
                _add_table(doc, headers, table_rows, aligns)
                continue

        # Bulleted / numbered lists
        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_markdown_runs(
                p, re.sub(r"^\s*[-*]\s+", "", line).strip()
            )
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_markdown_runs(
                p, re.sub(r"^\s*\d+\.\s+", "", line).strip()
            )
            i += 1
            continue

        # Blockquote
        if re.match(r"^\s*>\s*", line):
            p = doc.add_paragraph(style="Intense Quote")
            _add_inline_markdown_runs(p, re.sub(r"^\s*>\s*", "", line).strip())
            i += 1
            continue

        # Markdown image: ![caption](source)
        image_match = re.match(r"^\s*!\[(?P<caption>[^\]]*)\]\((?P<src>[^)]+)\)\s*$", stripped)
        if image_match:
            caption = image_match.group("caption").strip()
            src = image_match.group("src").strip()
            await _add_image_from_source(doc, src=src, caption=caption)
            i += 1
            continue

        # Horizontal rule
        if stripped in {"---", "***", "___"}:
            doc.add_paragraph("—" * 30)
            i += 1
            continue

        # Fallback: paragraph
        p = doc.add_paragraph()
        _add_inline_markdown_runs(p, stripped)
        i += 1

    if in_code_block and code_lines:
        _add_code_block(doc, "\n".join(code_lines))


def _set_header_footer(doc: Document, header_text: str, footer_text: str) -> None:
    for section in doc.sections:
        if header_text.strip():
            section.header.paragraphs[0].text = header_text.strip()
        if footer_text.strip():
            section.footer.paragraphs[0].text = footer_text.strip()


def _add_inline_markdown_runs(paragraph, text: str) -> None:  # noqa: ANN001
    # Supports: **bold**, *italic*, and `inline code`
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def _add_code_block(doc: Document, code_text: str) -> None:
    for line in (code_text.splitlines() or [""]):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def _add_image_placeholder(doc: Document, src: str, caption: str = "") -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"[image] {src}")
    run.italic = True
    run.font.size = Pt(9)
    if caption:
        cp = doc.add_paragraph(f"Figure: {caption}")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table_of_contents(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)

    toc_paragraph = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    hint = OxmlElement("w:t")
    hint.text = "Right-click and update field to generate TOC."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = toc_paragraph._p.add_r()
    r.append(fld_begin)
    r = toc_paragraph._p.add_r()
    r.append(instr_text)
    r = toc_paragraph._p.add_r()
    r.append(fld_separate)
    r = toc_paragraph._p.add_r()
    r.append(hint)
    r = toc_paragraph._p.add_r()
    r.append(fld_end)

    doc.add_paragraph("")


async def _append_images_to_doc(doc: Document, image_sources: str) -> None:
    items = [x.strip() for x in re.split(r"[\n,]", image_sources or "") if x.strip()]
    if not items:
        return
    for src in items:
        await _add_image_from_source(doc, src=src, caption=src)


async def _add_image_from_source(doc: Document, src: str, caption: str = "") -> None:
    image_bytes = await _load_image_bytes(src)
    if not image_bytes:
        _add_image_placeholder(doc, src=src, caption=caption or "Could not load image")
        return
    stream = io.BytesIO(image_bytes)
    try:
        doc.add_picture(stream)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:  # noqa: BLE001
        _add_image_placeholder(doc, src=src, caption=caption or "Unsupported image format")


async def _load_image_bytes(src: str) -> bytes | None:
    if src.startswith(("http://", "https://")):
        try:
            async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
                resp = await client.get(src)
                resp.raise_for_status()
                return resp.content
        except Exception:  # noqa: BLE001
            return None
    client = _get_minio()
    if client is None:
        return None
    try:
        return client.download_bytes(src)
    except Exception:  # noqa: BLE001
        return None


def _is_markdown_table_separator(line: str) -> bool:
    if "|" not in line:
        return False
    cells = _split_md_row(line)
    if not cells:
        return False
    for c in cells:
        token = c.replace(":", "").replace("-", "").strip()
        if token:
            return False
    return True


def _split_md_row(row: str) -> list[str]:
    trimmed = row.strip()
    if trimmed.startswith("|"):
        trimmed = trimmed[1:]
    if trimmed.endswith("|"):
        trimmed = trimmed[:-1]
    return [cell.strip() for cell in trimmed.split("|")]


def _parse_md_alignments(separator_row: str) -> list[str]:
    aligns: list[str] = []
    for cell in _split_md_row(separator_row):
        token = cell.strip()
        if token.startswith(":") and token.endswith(":"):
            aligns.append("center")
        elif token.endswith(":"):
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    aligns: list[str] | None = None,
) -> None:
    col_count = max(1, len(headers))
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    for idx in range(col_count):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        _add_inline_markdown_runs(cell.paragraphs[0], headers[idx] if idx < len(headers) else "")
        _align_cell(table.rows[0].cells[idx], aligns, idx)
    for row in rows:
        cells = table.add_row().cells
        for idx in range(col_count):
            cells[idx].text = ""
            _add_inline_markdown_runs(cells[idx].paragraphs[0], row[idx] if idx < len(row) else "")
            _align_cell(cells[idx], aligns, idx)


def _align_cell(cell, aligns: list[str] | None, idx: int) -> None:  # noqa: ANN001
    if not aligns or idx >= len(aligns):
        return
    align = aligns[idx]
    if align == "center":
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


async def _log_file_artifact(file_path: str, file_size: int, content_type: str) -> None:
    """Best-effort: insert a row in file_artifacts if a run context is active."""
    try:
        from agent_system.api.app import get_run_store
        from agent_system.core.run_context import get_run_context

        ctx = get_run_context()
        if ctx is None:
            return  # called outside an agent run — nothing to log
        store = get_run_store()
        if store is None:
            return
        await store.log_file_artifact(
            run_id=ctx.run_id,
            agent_name=ctx.agent_name,
            file_path=file_path,
            file_size=file_size,
            content_type=content_type,
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("write_file: could not log file_artifact: %s", exc)


@tool
def read_file(path: str) -> str:
    """Read a file from MinIO object storage and return its text content.

    Args:
        path: Object path / key inside the bucket, e.g. "reports/summary.txt".

    Returns:
        The file content as a string, or an error message.
    """
    client = _get_minio()
    if client is None:
        return "Error: MinIO storage is not available."
    try:
        raw = client.download_bytes(path)
        text = raw.decode("utf-8", errors="replace")
        if len(text) > 10000:
            text = text[:10000] + f"\n\n[... truncated — {len(raw)} bytes total]"
        return text
    except Exception as exc:  # noqa: BLE001
        return f"Error reading file '{path}': {exc}"


@tool
def list_files(prefix: str = "") -> str:
    """List files stored in MinIO under an optional path prefix.

    Args:
        prefix: Optional path prefix to filter results, e.g. "reports/".

    Returns:
        A newline-separated list of object paths, or a message if empty.
    """
    client = _get_minio()
    if client is None:
        return "Error: MinIO storage is not available."
    try:
        objects = client.list_objects(prefix=prefix)
        if not objects:
            label = f"under prefix '{prefix}'" if prefix else "in storage"
            return f"No files found {label}."
        return "\n".join(objects)
    except Exception as exc:  # noqa: BLE001
        return f"Error listing files: {exc}"


# ── Date / Time ───────────────────────────────────────────────────────────────

@tool
def get_datetime(timezone_name: str = "UTC") -> str:
    """Return the current date and time.

    Args:
        timezone_name: Timezone name — currently only "UTC" is supported.
                       For local time, pass "local".

    Returns:
        ISO 8601 formatted datetime string with timezone.
    """
    now = datetime.now(tz=timezone.utc)
    return now.strftime("%Y-%m-%d %H:%M:%S UTC (ISO: %Y-%m-%dT%H:%M:%SZ)")


# ── Text Utilities ────────────────────────────────────────────────────────────

@tool
def summarise_text(text: str, max_chars: int = 2000) -> str:
    """Truncate or excerpt a long text to a manageable length.

    Useful when passing large documents to the agent to avoid context overflow.

    Args:
        text: The full text to summarise / excerpt.
        max_chars: Maximum characters to keep (default 2 000).

    Returns:
        The truncated text with a note about total length if it was cut.
    """
    max_chars = max(100, min(max_chars, 20000))
    if len(text) <= max_chars:
        return text
    excerpt = text[:max_chars]
    # Try to break at a sentence boundary
    last_period = excerpt.rfind(". ", 0, max_chars)
    if last_period > max_chars // 2:
        excerpt = excerpt[: last_period + 1]
    return excerpt + f"\n\n[... {len(text) - len(excerpt)} more characters not shown]"


# ── Agent Memory (PostgreSQL-backed key-value store) ──────────────────────────

@tool
async def memory_save(agent_name: str, key: str, value: str) -> str:
    """Save a key-value pair to the agent's persistent memory in PostgreSQL.

    Use this to remember facts between separate runs (e.g. user preferences,
    intermediate results, or any state that must survive a conversation restart).

    Args:
        agent_name: The name of the agent saving the memory (use your own name).
        key: A short identifier for the memory (e.g. "user_city", "last_result").
        value: The value to store (always a string; use JSON for structured data).

    Returns:
        Confirmation message.
    """
    try:
        from agent_system.api.app import get_run_store
        await get_run_store().memory_save(agent_name=agent_name, key=key, value=value)
        logger.info("[memory_save] agent=%s key=%s", agent_name, key)
        return f"Memory saved: {key} = {value[:100]}"
    except Exception as exc:  # noqa: BLE001
        logger.error("[memory_save] failed: %s", exc)
        return f"Error saving memory: {exc}"


@tool
async def memory_get(agent_name: str, key: str) -> str:
    """Retrieve a value from the agent's persistent memory in PostgreSQL.

    Args:
        agent_name: The name of the agent whose memory to read.
        key: The key to look up.

    Returns:
        The stored value, or a message indicating the key was not found.
    """
    try:
        from agent_system.api.app import get_run_store
        value = await get_run_store().memory_get(agent_name=agent_name, key=key)
        if value is None:
            return f"No memory found for key '{key}'."
        logger.info("[memory_get] agent=%s key=%s found", agent_name, key)
        return value
    except Exception as exc:  # noqa: BLE001
        logger.error("[memory_get] failed: %s", exc)
        return f"Error reading memory: {exc}"


# ── Registry export ───────────────────────────────────────────────────────────

ALL_BUILTIN_TOOLS = [
    web_search,
    calculate,
    fetch_url,
    write_file,
    create_word_file,
    read_file,
    list_files,
    get_datetime,
    summarise_text,
    memory_save,
    memory_get,
]
